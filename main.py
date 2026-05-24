#!/usr/bin/env python3
"""
python main.py raw.txt [-o output.docx] [-c 1|2|3]
"""

import re
import os
import argparse
import math
from collections import OrderedDict
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from translate import Translator

translator = Translator(from_lang="en", to_lang="zh")

def auto_translate_word(raw_word):
    """直接翻译单词或短语，返回中文释义"""
    display_word = raw_word.replace('_', ' ')
    try:
        meaning = translator.translate(display_word)
    except:
        meaning = ''
    return meaning if meaning else ''

def parse_and_merge(filepath):
    """
    解析词汇文件，自动合并重复单词。
    规则：若有任意一行包含有效释义，则保留；若全都没有释义，则记为 None。
    返回 OrderedDict: {word: meaning_or_None}
    """
    merged = OrderedDict()
    with open(filepath, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            parts = re.split(r'\t| {2,}', line, maxsplit=1)
            word = parts[0].strip()
            meaning = None
            if len(parts) > 1:
                rest = parts[1].strip()
                pos_match = re.match(r'\(?([a-z]+\.)\)?\s*(.*)', rest, re.IGNORECASE)
                if pos_match:
                    meaning = pos_match.group(2).strip()
                else:
                    meaning = rest
                if not meaning:
                    meaning = None

            if meaning:
                merged[word] = meaning
            elif word not in merged:
                merged[word] = None
    return merged

def main():
    parser = argparse.ArgumentParser(description='Build printable vocabulary docx with multiple columns')
    parser.add_argument('raw_file')
    parser.add_argument('-o', '--output', default='vocabulary.docx')
    parser.add_argument('-c', '--columns', type=int, choices=[1, 2, 3], default=1,
                        help='单词并排列数(1~3, 默认 1)')
    args = parser.parse_args()

    # 合并与解析
    vocab = parse_and_merge(args.raw_file)
    unknown_count = sum(1 for v in vocab.values() if not v)
    known_count = len(vocab) - unknown_count
    print(f"已有翻译: {known_count} 条，未翻译: {unknown_count} 条")

    # 翻译所有无释义的单词
    for word in list(vocab.keys()):
        if not vocab[word]:
            meaning = auto_translate_word(word)
            vocab[word] = meaning
            print(f"\t自动处理: {word} -> {meaning}")

    # 排序
    sorted_words = sorted(vocab.keys(), key=lambda w: w.replace('_', ' ').lower())

    # ---------- 生成 docx ----------
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    title = doc.add_heading(f'Vocabulary List ({os.path.basename(args.raw_file)})', level=1)
    title.alignment = WD_TABLE_ALIGNMENT.CENTER

    cols = args.columns  # 单词并排列数
    total_cols = cols * 2  # 每对 (Word, Meaning) 占两栏
    rows = math.ceil(len(sorted_words) / cols)

    table = doc.add_table(rows=rows + 1, cols=total_cols, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hdr_cells = table.rows[0].cells
    for c in range(cols):
        hdr_cells[c * 2].text = 'Word'
        hdr_cells[c * 2 + 1].text = 'Meaning'
        for i in range(2):
            for p in hdr_cells[c * 2 + i].paragraphs:
                p.style.font.bold = True

    # 数据填充（行优先）
    for r in range(rows):
        row_cells = table.rows[r + 1].cells
        for c in range(cols):
            idx = r * cols + c
            if idx < len(sorted_words):
                word = sorted_words[idx]
                meaning = vocab[word]
                row_cells[c * 2].text = word.replace('_', ' ')
                row_cells[c * 2 + 1].text = meaning if meaning else ''
            else:
                # 留空
                pass

    # 列宽自动适应（简单方案，可根据需要微调）
    word_width = Inches(1.2)
    meaning_width = Inches(2.0)
    for c in range(cols):
        for cell in table.columns[c * 2].cells:
            cell.width = word_width
        for cell in table.columns[c * 2 + 1].cells:
            cell.width = meaning_width

    doc.save(args.output)
    print(f"✅ 词汇表已生成: {args.output}")

    # ---------- 输出回文本文件 ----------
    base, ext = os.path.splitext(args.raw_file)
    out_txt = f"{base}-output{ext}"
    with open(out_txt, 'w', encoding='utf-8') as f:
        for word in sorted_words:
            meaning = vocab[word]
            if meaning:
                f.write(f"{word}\t{meaning}\n")
            else:
                f.write(f"{word}\n")
    print(f"📄 完整词汇表已保存: {out_txt}")

if __name__ == '__main__':
    main()